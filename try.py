from docx import Document

def add_header_and_footer(doc, section_index):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml, OxmlElement, ns
    from docx.oxml.ns import nsdecls
    from docx.shared import Cm

    def create_element(name):
        return OxmlElement(name)

    def create_attribute(element, name, value):
        element.set(ns.qn(name), value)

    def add_page_number(run):
        fldChar1 = create_element('w:fldChar')
        create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = create_element('w:instrText')
        create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = create_element('w:fldChar')
        create_attribute(fldChar2, 'w:fldCharType', 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    
    # Set standard margins (1 inch on all sides)
    section = doc.sections[section_index]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    # Calculate content width (page width minus margins)
    page_width_inches = 8.27  # A4 width
    content_width_inches = page_width_inches - 2  # minus 2 inches for margins

    header = section.header
    footer = section.footer

    # Handle header
    for paragraph in header.paragraphs:
        paragraph.text = "Header"
        # Add bottom border using custom XML with #80c342 color
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        bottom = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="1" w:color="80C342"/></w:pBdr>')
        pPr.append(bottom)
        # Set minimal spacing after the text
        paragraph.paragraph_format.space_after = Pt(0)

    # Add empty paragraph for spacing in header
    spacing_para = header.add_paragraph()
    spacing_para.paragraph_format.space_before = Pt(6)
    spacing_para.paragraph_format.space_after = Pt(0)

    # Clear existing footer paragraphs
    footer.paragraphs.clear()

    # Create a table in the footer with smaller height and full content width
    table = footer.add_table(rows=1, cols=2, width=Inches(content_width_inches))
    # table.style = 'Table Grid'
    table.autofit = False
    
    # Set borders for cells - #80c342 top border, no other borders
    for row in table.rows:
        row.height = Inches(0.3)  # Set smaller row height
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="8" w:space="0" w:color="80C342"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
            tcPr.append(tcBorders)
    
    # Set column widths to half of content width
    half_width = content_width_inches / 2
    table.columns[0].width = Inches(half_width)
    table.columns[1].width = Inches(half_width)

    # Add footer text to left cell
    left_cell = table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.text = "Footer"
    left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Reduce paragraph spacing
    left_para.paragraph_format.space_before = Pt(0)
    left_para.paragraph_format.space_after = Pt(0)

    # Add page number to right cell
    right_cell = table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Reduce paragraph spacing
    right_para.paragraph_format.space_before = Pt(0)
    right_para.paragraph_format.space_after = Pt(0)
    run = right_para.add_run()
    add_page_number(run)

    return doc

if __name__ == "__main__":
    template = Document("template1.docx")
    template = add_header_and_footer(template, 1)

    template.save("template1_edited.docx")