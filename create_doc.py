def check_imports():
    try:
        import docx
    except ImportError:
        import subprocess
        import sys
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

def insert_element_before_paragraph(paragraph, element):
    """Insert a table after a specific paragraph."""
    # Get the paragraph element
    p_element = paragraph._element
    
    # Get the table element
    t_element = element._element
    
    # Insert the table after the paragraph
    p_element.addprevious(t_element)

def insert_element_after_paragraph(paragraph, element):
    """Insert a table after a specific paragraph."""
    # Get the paragraph element
    p_element = paragraph._element
    
    # Get the table element
    t_element = element._element
    
    # Insert the table after the paragraph
    p_element.addnext(t_element)

def find_marker_paragraph(doc, marker_text):
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == marker_text:
            return para
    return None

def add_element_at_marker(doc, element, marker_text):
    
    target_paragraph = find_marker_paragraph(doc, marker_text)

    if target_paragraph:
        # Move it to the desired position
        insert_element_before_paragraph(target_paragraph, element)

        return element
    else:
        raise ValueError("Marker paragraph not found")
        
def config():
    return {
        "MARKER_TEXT": "Start_new_section"
    }

def replace_paragraph_text(paragraph, old_text, new_text):
    from docx.shared import Pt
    inline = paragraph.runs
    # Loop added to work with runs (strings with same style)
    for i in range(len(inline)):
        if old_text in inline[i].text:
            text = inline[i].text.replace(old_text, new_text)
            inline[i].text = text
            if old_text == "[Title]" and len(new_text) > 8:
                inline[i].font.size = Pt(48)

def set_header_and_footer(doc, header_text, footer_text):
    from docx.shared import Pt
 
    header = doc.sections[1].header
    footer = doc.sections[1].footer

    header_para = header.paragraphs[0]
    footer_para = footer.paragraphs[0]

    replace_paragraph_text(header_para, "Header", header_text)
    replace_paragraph_text(footer_para, "Footer", footer_text)


    return doc

def add_clauses(doc, clauses_list: dict):
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import RGBColor
    import re
    # Use the same document as template if no template provided
    template_doc = doc


    clauses_list = sorted(clauses_list, key=lambda x: x['number'])

    # Find template paragraphs for each level
    heading_para = None
    clause_para = None
    subclause_para = None

    for paragraph in template_doc.paragraphs:
        if "Heading" in paragraph.text:
            heading_para = paragraph
        elif "Clause" in paragraph.text and "Subclause" not in paragraph.text:
            clause_para = paragraph
        elif "Subclause" in paragraph.text:
            subclause_para = paragraph

    if not all([heading_para, clause_para, subclause_para]):
        raise ValueError("Could not find all required template paragraphs (Heading, Clause, Subclause)")

    original_numPr = heading_para._p.pPr.numPr

    def create_num_pr(ilvl_val):
        """Create numbering properties for the specified level."""
        num_pr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), str(ilvl_val))
        numId = OxmlElement('w:numId')
        numId.set(qn('w:val'), str(original_numPr.numId.val))
        num_pr.append(ilvl)
        num_pr.append(numId)
        return num_pr

    def copy_font_properties(source_run, target_run):
        """Copy font properties from source run to target run."""
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        if hasattr(source_run.font, 'color') and source_run.font.color:
            target_run.font.color.rgb = source_run.font.color.rgb

    def make_xml_compatible(s):
        # Find all non-XML-compatible control characters
        removed = [f"{ord(c):02x}" for c in s if ord(c) < 32 and c not in (9, 10, 13)]
        # if removed:
        #     print(f"Removed control characters (hex): {removed}")
        # Remove all control characters except tab, newline, carriage return
        return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', s)

    # Add each clause with proper formatting
    for clause in clauses_list:
        # Add clause title (level 1)
        p1 = doc.add_paragraph()
        p1.style = heading_para.style
        p1._p.get_or_add_pPr().append(create_num_pr(0))
        run1 = p1.add_run(f"{clause['name']}")
        copy_font_properties(heading_para.runs[0], run1)

        add_element_at_marker(doc, p1, "Marker_paragraph_for_part_2")

        # Add each clause content (level 2)
        for content in clause['contents']:
            cleaned_content = make_xml_compatible(content['content'])
            _split = cleaned_content.split('~~')
            _split = [chunk for chunk in _split if chunk.strip()]
            __split = [chunk.split('~') for chunk in _split]
            all_chunks = []
            for chunk in __split:
                for subchunk in chunk:
                    if subchunk.strip():
                        all_chunks.append(subchunk)
                
            p2 = doc.add_paragraph()
            p2.style = clause_para.style
            p2._p.get_or_add_pPr().append(create_num_pr(1))
            try:
                for i, chunk in enumerate(all_chunks):
                    run2 = p2.add_run(chunk)
                    copy_font_properties(clause_para.runs[0], run2)
                    if i == 1:
                        run2.font.color.rgb = RGBColor(255, 0, 0)
                        run2.font.strike = True
                    elif i > 1:
                        run2.font.color.rgb = RGBColor(0, 0, 255)
            except Exception as e:
                print(content['content'])
                raise ValueError(f"Error adding clause content: {e}")

            add_element_at_marker(doc, p2, "Marker_paragraph_for_part_2")

            # Add subtopics (level 3)
            for subtopic in content['subtopics']:
                cleaned_subtopic = make_xml_compatible(subtopic)
                _split = cleaned_subtopic.split('~~')
                _split = [chunk for chunk in _split if chunk.strip()]
                __split = [chunk.split('~') for chunk in _split]
                all_chunks = []
                for chunk in __split:
                    for subchunk in chunk:
                        if subchunk.strip():
                            all_chunks.append(subchunk)

                p3 = doc.add_paragraph()
                p3.style = subtopic_para.style
                p3._p.get_or_add_pPr().append(create_num_pr(2))
                try:
                    for i, chunk in enumerate(all_chunks):
                        run3 = p3.add_run(chunk)
                        copy_font_properties(subtopic_para.runs[0], run3)
                        if i == 1:
                            run3.font.color.rgb = RGBColor(255, 0, 0)
                            run3.font.strike = True
                        elif i > 1:
                            run3.font.color.rgb = RGBColor(0, 0, 255)
                except Exception as e:
                    print(subtopic)
                    raise ValueError(f"Error adding subtopic content: {e}")

                add_element_at_marker(doc, p3, "Marker_paragraph_for_part_2")

    heading_para._element.getparent().remove(heading_para._element)
    clause_para._element.getparent().remove(clause_para._element)
    subclause_para._element.getparent().remove(subclause_para._element)

    marker_paragraph = find_marker_paragraph(doc, "Marker_paragraph_for_part_2")
    marker_paragraph._element.getparent().remove(marker_paragraph._element)

    return doc

def set_updatefields_true(doc):
    import lxml
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    # add child to doc.settings element
    element_updatefields = lxml.etree.SubElement(
        doc.settings.element, f"{namespace}updateFields"
    )
    element_updatefields.set(f"{namespace}val", "true")
    return doc

def add_part_i(doc, list_of_rows: dict, marker_text):
    """
    rows: list[dict]
    keys: query, answer, answer_type
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.section import WD_SECTION
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    
    # Create individual tables for each row
    from docx.shared import Pt
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    for i, row_data in enumerate(list_of_rows):
        # Determine which columns to include based on non-empty data
        query = row_data.get('query', '').strip()
        answer = row_data.get('answer', '').strip()
        answer_type = row_data.get('answer_type', '').strip()
        
        # Build list of columns to include
        columns_data = []
        column_alignments = []
        
        is_query = False
        if query:
            columns_data.append(query)
            column_alignments.append(WD_ALIGN_PARAGRAPH.LEFT)
            is_query = True

        if answer:
            columns_data.append(answer)
            column_alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
        
        if answer_type:
            columns_data.append(answer_type)
            column_alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
        
        # Skip if no data to display
        if not columns_data:
            continue
            
        # Create table with dynamic number of columns
        table = doc.add_table(rows=1, cols=len(columns_data))
        add_element_at_marker(doc, table, marker_text)
        # table.style = 'Table Grid'
        
        # Customize table borders
        tbl = table._tbl
        tblPr = tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')
        
        
        for border_name in ['bottom', 'insideH']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '8')  # 1px = 8 eighths of a point
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'b8b9bb')  # rgb(184, 185, 187) in hex
            tblBorders.append(border)
        
        # Remove vertical borders (left, right, insideV)
        for border_name in ['left', 'right', 'insideV', 'top']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
            tblBorders.append(border)
        
        
        tblPr.append(tblBorders)
        
        row = table.rows[0]
        
        # Set cell margins (10px top and bottom padding)
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcMar = OxmlElement('w:tcMar')
            
            # Set top margin
            top = OxmlElement('w:top')
            top.set(qn('w:w'), '150')  # 10px in twentieths of a point
            top.set(qn('w:type'), 'dxa')
            tcMar.append(top)
            
            # Set bottom margin
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:w'), '150')  # 10px in twentieths of a point
            bottom.set(qn('w:type'), 'dxa')
            tcMar.append(bottom)
            
            tcPr.append(tcMar)
        
        # Fill cells with data and apply formatting
        for j, (data, alignment) in enumerate(zip(columns_data, column_alignments)):
            cell = row.cells[j]
            cell.text = data
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            for paragraph in cell.paragraphs:
                paragraph.alignment = alignment
                for run in paragraph.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    if is_query and j == 0: 
                        run.font.bold = True
    
    marker_paragraph = find_marker_paragraph(doc, marker_text)
    marker_paragraph._element.getparent().remove(marker_paragraph._element)

    return doc

def add_title_page(doc, title, subtitle, date, additional_information):
    from docx.enum.section import WD_SECTION
    from docx import Document
    from docx.shared import Pt, RGBColor
    title_page = doc

    # Get the paragraphs
    title_paragraph = title_page.paragraphs[3]
    subtitle_paragraph = title_page.paragraphs[4]
    date_paragraph = title_page.paragraphs[6]
    additional_information_paragraph = title_page.paragraphs[12]

    # Replace text while preserving styles
    replace_paragraph_text(title_paragraph, "[Title]", title)
    replace_paragraph_text(subtitle_paragraph, "[Subtitle]", subtitle)
    replace_paragraph_text(date_paragraph, "[Date]", date)
    replace_paragraph_text(additional_information_paragraph, "[Additional Information 1]", additional_information)

    return doc

def download_file_from_google_drive(file_id: str, destination: str):
    import requests
    """
    Downloads a file from Google Drive given its file ID and saves it to the specified destination.
    The file must be accessible to everyone with the link.
    Args:
        file_id (str): The Google Drive file ID.
        destination (str): The local path to save the downloaded file.
    """
    URL = "https://drive.usercontent.google.com/u/0/uc"
    params = {"id": file_id, "export": "download"}
    response = requests.get(URL, params=params, stream=True)
    response.raise_for_status()
    with open(destination, "wb") as f:
        print(f"Downloading file to {destination}")
        for chunk in response.iter_content(chunk_size=8192):
            if chunk:
                f.write(chunk)
        f.close()

def create_docx_file(input_file_name, output_file_name, clauses: dict, preamble_rows: dict):
    from docx import Document
    doc = Document(input_file_name)
    doc = add_title_page(doc, "BPVOY5", "Voyage Charter Party", "February 2016", "This is the additional information")
    doc = set_header_and_footer(doc, "Header of my document", "Footer of my document")

    doc = add_part_i(doc, preamble_rows, "Marker_paragraph_for_preamble")
    doc = add_part_i(doc, preamble_rows, "Marker_paragraph_for_part_1")

    doc = add_clauses(doc, clauses)
    doc = set_updatefields_true(doc)

    doc.save("chat_id_" + output_file_name) # TODO: replace with actual chat_id implementation

def get_real_filename(filename):
    """
    USED ONLY IN LOCAL TESTING
    TODO: remove this function in production
    """
    return "2760eadc_806e_42c6_b9a3_ac92a56f1aeb_" + filename

def read_file(filename):
    """
    USED ONLY IN LOCAL TESTING
    TODO: remove this function in production
    """
    import os
    import json
    format = filename.split(".")[-1]
    real_filename = get_real_filename(filename)
    if not os.path.exists(real_filename):
        raise FileNotFoundError(f"File {real_filename} not found")
    if "json" in format:
        with open(real_filename, "r", encoding="utf-8") as file:
            return json.load(file)
        file.close()
    else:
        with open(real_filename, "r", encoding="utf-8") as file:
            return file.read()
        file.close()

def create_docx_file(flat_history: dict):
    """
    TODO: implement the storage of the **template** docx file in the database
    TODO: implement the storage of the output docx file in the database
    TODO: implement sorting of all topics
    TODO: implement strikethrough, blue/red/black text, etc.
    TODO: implement correct numbering of clause paragraphs
    """
    # from parse_flat_history import parse_flat_history
    parsed_flat_history = parse_flat_history(flat_history)

    check_imports()
    from docx import Document
    real_template_file_name = get_real_filename("perfect_template.docx")
    download_file_from_google_drive("1Np_RB7uvQxViHzGwzl0IK_XKz15UA-f8", real_template_file_name)
    doc = Document(real_template_file_name)
    doc_data = parsed_flat_history['title_page']
    doc = add_title_page(doc, doc_data["doc_title"], doc_data["doc_subtitle"], doc_data["date"], doc_data["additional_info"])
    doc = set_header_and_footer(doc, doc_data["header"], doc_data["footer"])
    part_i_rows = parsed_flat_history['part_i']
    preamble_rows = parsed_flat_history['preamble']
    clauses = parsed_flat_history['part_ii']
    doc = add_part_i(doc, preamble_rows, "Marker_paragraph_for_preamble")
    doc = add_part_i(doc, part_i_rows, "Marker_paragraph_for_part_1")
    doc = add_clauses(doc, clauses)
    doc = set_updatefields_true(doc)
    doc.save(get_real_filename(f"output.docx"))

    upload_file_asset(
        file_name=get_real_filename(f"output.docx"),
        file_bytes=open(get_real_filename(f"output.docx"), "rb").read(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    delete_file("output.docx")
    delete_file("perfect_template.docx")

if __name__ == "__main__":
    import json
    full_flat_history = None
    with open('full_flat_history.json', 'r') as f:
        full_flat_history = json.load(f)

    create_docx_file(full_flat_history)