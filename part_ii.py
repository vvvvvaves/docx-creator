def add_heading(doc, heading_text):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_LINE_SPACING

    # Create and configure custom heading style
    custom_style = doc.styles['Title']
    custom_style.font.size = Pt(24)
    custom_style.font.name = 'Arial'
    custom_style.font.color.rgb = RGBColor(0, 166, 80)
    custom_style.paragraph_format.space_after = Pt(0)
    custom_style.paragraph_format.space_before = Pt(0)

    custom_style.paragraph_format._element.xpath('.//w:pBdr')[0].getparent().remove(
        custom_style.paragraph_format._element.xpath('.//w:pBdr')[0])

    # Remove empty paragraphs at the start of the document
    if doc.paragraphs and not doc.paragraphs[0].text.strip():
        p = doc.paragraphs[0]._element
        p.getparent().remove(p)
        p._p = p._element = None

    # Add heading using the custom style
    heading = doc.add_heading(heading_text, 0)
    heading.style = custom_style

    return doc

def add_clause(doc):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_LINE_SPACING, WD_TAB_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Twips

    # Create paragraph style for the list
    style = doc.styles.add_style('Clause 1', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0, 176, 80)  # Standard green

    style = doc.styles.add_style('Clause 2', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    indent_coef = 1.26
    style.paragraph_format.first_line_indent = Cm(-indent_coef*1)
    style.paragraph_format.left_indent = Cm(indent_coef*2)  # Set left indent to 0.63 cm

    style = doc.styles.add_style('Clause 3', WD_STYLE_TYPE.PARAGRAPH)
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.first_line_indent = Cm(indent_coef*(-0.5))
    style.paragraph_format.left_indent = Cm(indent_coef*2.5)

    # Add the list item with custom numbering
    paragraph = doc.add_paragraph(style='Clause 1')
    run = paragraph.add_run("1\t")  # Add number
    run = paragraph.add_run("Lorem Ipsum Dolor")  # Add tab and text

    paragraph_2 = doc.add_paragraph(style='Clause 2')
    run = paragraph_2.add_run("1.1\t")  # Add number
    run = paragraph_2.add_run("Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua:")  # Add tab and text

    paragraph_3 = doc.add_paragraph(style='Clause 3')
    run = paragraph_3.add_run("a)\t") # Add letter
    run = paragraph_3.add_run("Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat. Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur.")  # Add tab and text

    return doc

if __name__ == "__main__":
    from docx import Document
    doc = Document("header_and_footer_empty.docx")
    doc = add_heading(doc, "Part 2")
    doc = add_clause(doc)
    doc.save("part_ii.docx")